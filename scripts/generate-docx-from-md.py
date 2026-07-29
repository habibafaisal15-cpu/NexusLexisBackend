"""Convert a markdown file in docs/ to a Word document with header/footer."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_header_footer(doc: Document, meta: dict[str, str]) -> None:
    title = meta.get("Document Title", "Nexus Lexis Document")
    doc_id = meta.get("Document ID", "NL-DOC")
    version = meta.get("Version", "1.0")
    doc_class = meta.get("Classification", "Internal")
    doc_date = meta.get("Last Updated", "")

    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    hp = section.header.paragraphs[0]
    hp.text = f"{title}  |  {doc_id}  |  v{version}  |  {doc_class}"
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in hp.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    fp = section.footer.paragraphs[0]
    fp.text = f"© Nexus Lexis · nexuslexis.law  |  Updated: {doc_date}  |  Page "
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in fp.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run = fp.add_run()
    for el in (
        ("begin", None),
        (None, " PAGE "),
        ("separate", None),
        ("end", None),
    ):
        if el[1]:
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = el[1]
            run._r.append(instr)
        else:
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), el[0])
            run._r.append(fld)
    run.font.size = Pt(8)


def add_title_page(doc: Document, meta: dict[str, str]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("NEXUS LEXIS PLATFORM")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x16, 0x21, 0x3E)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = meta.get("Document Title", "API Reference")
    r = p.add_run(title.split("—")[-1].strip() if "—" in title else title)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x16, 0x21, 0x3E)

    doc.add_paragraph()
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for key in ("Document ID", "Version", "Last Updated", "Classification", "Owner", "Applies To"):
        if key in meta:
            row = table.add_row().cells
            row[0].text = key
            row[1].text = meta[key]
            set_cell_shading(row[0], "E8ECF4")
    doc.add_page_break()


def add_rich_text(paragraph, text: str) -> None:
    for part in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(cols):
            cell = table.rows[ri].cells[ci]
            text = row[ci] if ci < len(row) else ""
            cell.text = ""
            para = cell.paragraphs[0]
            add_rich_text(para, text.strip())
            for run in para.runs:
                run.font.size = Pt(8)
            if ri == 0:
                set_cell_shading(cell, "D4E4F7")
                for run in para.runs:
                    run.bold = True


def convert(md_text: str) -> Document:
    meta: dict[str, str] = {}
    body: list[str] = []
    in_meta = False
    for line in md_text.splitlines():
        if line.strip() == "---":
            in_meta = not in_meta
            continue
        if in_meta and ":" in line:
            k, v = line.split(":", 1)
            meta[k.strip()] = v.strip()
        elif not in_meta:
            body.append(line)

    doc = Document()
    add_header_footer(doc, meta)
    add_title_page(doc, meta)

    in_code = False
    code_lines: list[str] = []
    table_rows: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_rows
        if table_rows:
            add_table(doc, table_rows)
            table_rows = []
            doc.add_paragraph()

    for line in body:
        if in_code:
            if line.strip().startswith("```"):
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(8)
                in_code = False
                code_lines = []
            else:
                code_lines.append(line)
            continue

        if line.strip().startswith("```"):
            flush_table()
            in_code = True
            continue

        if line.strip().startswith("|") and "|" in line.strip()[1:]:
            if re.match(r"^\|\s*[-: ]+\|", line.strip()):
                continue
            table_rows.append([c.strip() for c in line.strip().strip("|").split("|")])
            continue

        flush_table()
        if line.strip() == "---":
            doc.add_paragraph()
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), 1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), 2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), 3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), 4)
        elif line.startswith("> "):
            p = doc.add_paragraph(style="Intense Quote")
            add_rich_text(p, line[2:])
        elif line.strip():
            p = doc.add_paragraph()
            add_rich_text(p, line.strip())

    flush_table()
    return doc


def main() -> None:
    stem = sys.argv[1] if len(sys.argv) > 1 else "NEXUSLEXIS-API-REFERENCE"
    md_path = ROOT / "docs" / f"{stem}.md"
    out_path = ROOT / "docs" / f"{stem}.docx"
    doc = convert(md_path.read_text(encoding="utf-8"))
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Created: {out_path}")


if __name__ == "__main__":
    main()
