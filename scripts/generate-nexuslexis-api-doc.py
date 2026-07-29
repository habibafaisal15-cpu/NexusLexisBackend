"""Generate Nexus Lexis API Documentation Word doc (Brilliance template style)."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "NexusLexis_API_Documentation.docx"


def mono(doc: Document, text: str, bold: bool = False, size: int = 10) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(size)
    run.bold = bold


def body(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.bold = bold


def section(doc: Document, title: str) -> None:
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(12)


def endpoint(
    doc: Document,
    method: str,
    path: str,
    *,
    auth: str = "No auth",
    roles: str = "",
    headers: str = "",
    query: str = "",
    body: str = "",
    note: str = "",
) -> None:
    doc.add_paragraph()
    mono(doc, f"{method} {path}", bold=True)
    mono(doc, f"Auth: {auth}")
    if roles:
        mono(doc, f"Roles: {roles}")
    if headers:
        mono(doc, f"Headers: {headers}")
    if query:
        mono(doc, f"Query: {query}")
    if body:
        for line in body.split("\n"):
            mono(doc, line)
    if note:
        mono(doc, f"Note: {note}")


def build_document() -> Document:
    doc = Document()

    # ── Cover / intro (Brilliance style) ─────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Nexus Lexis Platform")
    r.bold = True
    r.font.size = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("API Documentation")
    r.bold = True
    r.font.size = Pt(16)

    doc.add_paragraph()
    intro = """Main API Base URL: http://localhost:3000
Auth API Base URL: http://localhost:3001
  (All Auth routes also available via proxy: http://localhost:3000/api/auth/*)
LEX AI Base URL: http://localhost:8001
  (LEX routes also proxied: http://localhost:3000/api/v1/lex/*)

Stack: Node.js + Express + PostgreSQL + Django (LEX AI) + JWT + WebSocket
All protected routes need header: Authorization: Bearer <token>
Optional dev headers (Main API): X-Client-Role, X-Workspace-Context, ngrok-skip-browser-warning

Roles: client | lawyer | ca | admin
(JWT may also use labels: CorporateClient | LegalAdvocate | CharteredAccountant | Admin)

Typical auth success: { token, user } or { profile, token }
Typical error: { error: "message" } or { ok: false, code, error }
OTP endpoints return: { ok, message, expiresInMinutes } or { ok, verificationToken }

Database: nexuslexis (PostgreSQL)
LEX AI local DB: lex_backend/db.sqlite3 (SQLite, dev)

Mainsite signup flow:
1. POST /api/auth/register/send-otp with email
2. POST /api/auth/register/verify-otp → receive verificationToken
3. POST /api/auth/register with verificationToken + profile fields
4. PUT /api/auth/profile/client to complete client profile
5. POST /api/auth/profile/switch-role to open role dashboard

Source files: server.js · auth_backend/routes/*.js · routes/*.js · lex_backend/lex_ai/"""
    for line in intro.split("\n"):
        mono(doc, line)

    # ── 0. Health ────────────────────────────────────────────────────────────
    section(doc, "0. Health")
    endpoint(doc, "GET", "/api/health", auth="No auth", note="Main API health + DB status. Source: server.js")
    endpoint(doc, "GET", "/api/health", auth="No auth", note="Auth API health. Base :3001. Source: auth_backend/server.js")
    endpoint(doc, "GET", "/", auth="No auth", note="Auth API service index + endpoint list. Base :3001")

    # ── 1. Auth /api/auth ────────────────────────────────────────────────────
    section(doc, "1. Auth /api/auth")
    body(doc, "Base :3001 or proxy http://localhost:3000/api/auth. Source: auth_backend/routes/auth.js")
    body(doc, "verificationToken required on register (from verify-otp). role: client | lawyer | ca")

    endpoint(
        doc, "POST", "/api/auth/register/send-otp", auth="No auth",
        headers="Content-Type: application/json",
        body='Body (JSON):\n{\n  "email": "user@example.com"\n}',
        note="Sends 6-digit OTP email (10 min TTL). Returns { ok, email, message, expiresInMinutes }. Rate limit: 5 sends / 15 min.",
    )
    endpoint(
        doc, "POST", "/api/auth/register/verify-otp", auth="No auth",
        headers="Content-Type: application/json",
        body='Body (JSON):\n{\n  "email": "user@example.com",\n  "code": "482915"\n}',
        note='Alias: "otp" accepted instead of "code". Returns { ok, verificationToken }.',
    )
    endpoint(
        doc, "POST", "/api/auth/register/validate", auth="No auth",
        headers="Content-Type: application/json",
        body='Body (JSON):\n{\n  "email": "user@example.com"\n}',
        note="Returns { valid, available } before signup.",
    )
    endpoint(
        doc, "POST", "/api/auth/register", auth="No auth",
        headers="Content-Type: application/json",
        body='Body (JSON):\n{\n  "fullName": "Ali Khan",\n  "email": "ali@example.com",\n  "password": "secret123",\n  "phone": "03001234567",\n  "role": "client",\n  "verificationToken": "<from verify-otp>"\n}',
        note="Creates account. Returns { token, user }. role defaults to client.",
    )
    endpoint(
        doc, "POST", "/api/auth/login", auth="No auth",
        headers="Content-Type: application/json",
        body='Body (JSON):\n{\n  "email": "ali@example.com",\n  "password": "secret123"\n}',
        note="Returns { token, user } with profile bundle when dashboard user exists.",
    )
    endpoint(
        doc, "GET", "/api/auth/me", auth="Bearer JWT required",
        headers="Authorization: Bearer <token>",
        note="Returns current user + full profile (roles, verification status, canApply flags).",
    )
    endpoint(
        doc, "GET", "/api/auth/roles", auth="No auth",
        note='Returns { registerRoles: ["client","lawyer","ca"], allRoles }.',
    )
    endpoint(
        doc, "GET", "/api/auth/google/url", auth="No auth",
        query="?state=login",
        note="Returns { url } for Google OAuth redirect. Requires GOOGLE_CLIENT_ID in env.",
    )
    endpoint(
        doc, "GET", "/api/auth/google/callback", auth="No auth",
        query="?code=<auth_code>&state=login",
        note="OAuth callback. Redirects to frontend /login?token=...",
    )
    endpoint(
        doc, "POST", "/api/auth/google/token", auth="No auth",
        headers="Content-Type: application/json",
        body='Body (JSON):\n{\n  "idToken": "<google credential>",\n  "role": "client"\n}',
        note='Accepts "credential" alias for idToken. Returns { token, user }.',
    )

    # ── 2. Profile /api/auth/profile ─────────────────────────────────────────
    section(doc, "2. Profile /api/auth/profile")
    body(doc, "Source: auth_backend/routes/profile.js")

    endpoint(doc, "GET", "/api/auth/profile", auth="Bearer JWT required", headers="Authorization: Bearer <token>", note="Full profile bundle.")
    endpoint(
        doc, "PUT", "/api/auth/profile/client", auth="Bearer JWT required",
        headers="Content-Type: application/json | Authorization: Bearer <token>",
        body='Body (JSON):\n{\n  "cnic": "35201-1234567-1",\n  "address": "123 Main St",\n  "city": "Lahore",\n  "profilePhoto": "<url or filename>",\n  "phone": "03001234567"\n}',
        note="Saves client profile. Returns { profile, token } with refreshed JWT.",
    )
    endpoint(
        doc, "POST", "/api/auth/profile/client/signup", auth="Bearer JWT required",
        headers="Authorization: Bearer <token>",
        note="Complete client signup flow shortcut. Returns { profile, token }.",
    )
    endpoint(
        doc, "POST", "/api/auth/profile/lawyer/apply", auth="Bearer JWT required",
        headers="Content-Type: application/json | Authorization: Bearer <token>",
        body='Body (JSON):\n{\n  "fullName": "Adv. Ali Khan",\n  "profilePhoto": "...",\n  "barCertificate": "...",\n  "cnicFront": "...",\n  "cnicBack": "...",\n  "barCouncilName": "Punjab Bar Council",\n  "barCouncilNumber": "L-12345",\n  "city": "Lahore",\n  "practiceAreas": ["Family Law"],\n  "languages": ["English", "Urdu"],\n  "shortBio": "...",\n  "fullBio": "...",\n  "officeAddress": "...",\n  "onlineFee": 5000,\n  "inPersonFee": 8000,\n  "consultationMode": "Both"\n}',
        note="Submits lawyer verification application. Status → Pending. Returns { profile, token }.",
    )
    endpoint(
        doc, "POST", "/api/auth/profile/ca/apply", auth="Bearer JWT required",
        headers="Content-Type: application/json | Authorization: Bearer <token>",
        body='Body (JSON):\n{\n  "fullName": "CA Sara Ahmed",\n  "photo": "...",\n  "caCertificate": "...",\n  "cnicFront": "...",\n  "cnicBack": "...",\n  "cnic": "35201-1234567-1",\n  "icapMembershipNo": "ICAP-1234",\n  "qualification": "FCA",\n  "city": "Karachi",\n  "serviceAreas": ["FBR Income Tax Filing"],\n  "fees": 15000,\n  "availability": "Mon-Fri 9-5"\n}',
        note="Submits CA verification application. Returns { profile, token }.",
    )
    endpoint(
        doc, "POST", "/api/auth/profile/switch-role", auth="Bearer JWT required",
        headers="Content-Type: application/json | Authorization: Bearer <token>",
        body='Body (JSON):\n{\n  "role": "lawyer"\n}',
        note='role: client | lawyer | ca. Returns 409 PENDING_VERIFICATION if role not approved. Returns { profile, token }.',
    )

    # ── 3. Admin /api/auth/admin ─────────────────────────────────────────────
    section(doc, "3. Admin /api/auth/admin")
    body(doc, "Source: auth_backend/routes/admin.js. Requires admin role in JWT.")

    endpoint(doc, "GET", "/api/auth/admin/applications", auth="Bearer JWT + admin", headers="Authorization: Bearer <token>", note="List pending lawyer/CA applications.")
    endpoint(doc, "GET", "/api/auth/admin/applications/:userId", auth="Bearer JWT + admin", query="?type=lawyer | ?type=ca", note="Application details + uploaded documents.")
    endpoint(
        doc, "POST", "/api/auth/admin/applications/:userId/approve", auth="Bearer JWT + admin",
        headers="Content-Type: application/json | Authorization: Bearer <token>",
        body='Body (JSON):\n{\n  "type": "lawyer"\n}',
        note="Approves application. type: lawyer | ca.",
    )
    endpoint(
        doc, "POST", "/api/auth/admin/applications/:userId/reject", auth="Bearer JWT + admin",
        headers="Content-Type: application/json | Authorization: Bearer <token>",
        body='Body (JSON):\n{\n  "type": "ca"\n}',
        note="Rejects application. User may re-apply.",
    )

    # ── 4. Legacy Client Portal Auth /api/v2/auth ────────────────────────────
    section(doc, "4. Legacy Client Portal Auth /api/v2/auth (Main API local)")
    body(doc, "Source: server.js. Corporate client portal only — prefer /api/auth for mainsite.")

    endpoint(
        doc, "POST", "/api/v2/auth/register", auth="No auth",
        headers="Content-Type: application/json",
        body='Body (JSON):\n{\n  "name": "Habib Corp",\n  "email": "corp@example.com",\n  "password": "secret123"\n}',
        note="Registers corporate client in Main API users table. Returns { token, user }.",
    )
    endpoint(
        doc, "POST", "/api/v2/auth/login", auth="No auth",
        headers="Content-Type: application/json",
        body='Body (JSON):\n{\n  "email": "corp@example.com",\n  "password": "secret123"\n}',
        note="403 if account role is not client.",
    )
    endpoint(doc, "GET", "/api/v2/auth/me", auth="Bearer JWT required", headers="Authorization: Bearer <token>", note="Current client user.")
    endpoint(doc, "POST", "/api/v2/auth/session", auth="No auth", note="Demo client session token (habibcorp@nexuslexis.law).")

    # ── 5. Client Workspace /api/v2 ──────────────────────────────────────────
    section(doc, "5. Client Workspace /api/v2")
    body(doc, "Source: server.js. Auth: Bearer JWT (client session).")

    endpoint(doc, "GET", "/api/v2/workspace", auth="Bearer JWT required", note="Bootstrap all dashboard data for client.")
    endpoint(doc, "GET", "/api/v2/notifications", auth="Bearer JWT required", note="List client notifications.")
    endpoint(doc, "DELETE", "/api/v2/notifications/:id", auth="Bearer JWT required", note="Dismiss one notification.")
    endpoint(doc, "DELETE", "/api/v2/notifications", auth="Bearer JWT required", note="Clear all notifications.")

    # ── 6. Document Library ──────────────────────────────────────────────────
    section(doc, "6. Document Library /api/v2/library")
    endpoint(doc, "GET", "/api/v2/library/catalog", auth="No auth", query="?category=&search=", note="Public template catalog.")
    endpoint(doc, "GET", "/api/v2/library/templates/:slug", auth="No auth", note="Single template by slug.")

    # ── 7. Client Documents & Orders ─────────────────────────────────────────
    section(doc, "7. Client Documents & Orders /api/v2")
    endpoint(doc, "GET", "/api/v2/documents", auth="Bearer JWT required", query="?status=", note="List client document orders.")
    endpoint(doc, "GET", "/api/v2/documents/:orderNumber", auth="Bearer JWT required", note="Single document order detail.")
    endpoint(doc, "GET", "/api/v2/documents/:orderNumber/download", auth="Bearer JWT required", note="Download completed document file.")
    endpoint(doc, "GET", "/api/v2/orders", auth="Bearer JWT required", note="List orders.")
    endpoint(
        doc, "POST", "/api/v2/orders", auth="Bearer JWT required",
        headers="Content-Type: application/json | Authorization: Bearer <token>",
        body='Body (JSON):\n{\n  "templateId": "nda-template",\n  "templateName": "NDA",\n  "formData": { "partyA": "..." }\n}',
        note="Create document order from library template.",
    )

    # ── 8. VLO Matters ───────────────────────────────────────────────────────
    section(doc, "8. Virtual Legal Office /api/v2/vlo")
    endpoint(doc, "GET", "/api/v2/vlo/matters", auth="Bearer JWT required", note="List client VLO matters.")
    endpoint(
        doc, "POST", "/api/v2/vlo/matters", auth="Bearer JWT required",
        headers="Content-Type: multipart/form-data | Authorization: Bearer <token>",
        body='Body (multipart/form-data):\ntitle: Contract review\ndescription: Need opinion on vendor agreement\nfiles: <up to 10 files>',
        note="Submit new legal matter with attachments.",
    )
    endpoint(doc, "GET", "/api/vlo/matters/download/:id", auth="Bearer JWT required", note="Download advisory opinion PDF (legacy path).")

    # ── 9. Appointments & Subscription ───────────────────────────────────────
    section(doc, "9. Appointments, Subscription & Billing /api/v2")
    endpoint(
        doc, "POST", "/api/v2/appointments", auth="Bearer JWT required",
        headers="Content-Type: application/json | Authorization: Bearer <token>",
        body='Body (JSON) — Lawyer booking:\n{\n  "lawyerProfileId": 1,\n  "lawyerName": "Adv. Ali",\n  "slot": "2026-08-01T10:00:00",\n  "mode": "online",\n  "intake": "Need family law advice",\n  "clientCity": "Lahore"\n}\n\nBody (JSON) — CA booking:\n{\n  "professionalRole": "CA",\n  "caProfileId": 2,\n  "caName": "CA Sara",\n  "slot": "2026-08-01T14:00:00",\n  "mode": "in-person",\n  "intake": "Tax filing help"\n}',
        note="Books lawyer or CA appointment.",
    )
    endpoint(doc, "GET", "/api/v2/appointments", auth="Bearer JWT required", note="List client appointments.")
    endpoint(doc, "GET", "/api/v2/subscription", auth="Bearer JWT required", note="Client subscription details.")
    endpoint(doc, "POST", "/api/v2/subscription/cancel", auth="Bearer JWT required", note="Cancel subscription. Returns { success: true }.")
    endpoint(doc, "GET", "/api/v2/invoices", auth="Bearer JWT required", note="List client invoices.")
    endpoint(
        doc, "POST", "/api/v2/evaluations", auth="Bearer JWT required",
        headers="Content-Type: application/json | Authorization: Bearer <token>",
        body='Body (JSON):\n{\n  "rating": 5,\n  "comment": "Excellent service",\n  "threadId": "thread-123"\n}',
        note="Submit service evaluation after consultation.",
    )

    # ── 10. Directory ────────────────────────────────────────────────────────
    section(doc, "10. Lawyer & CA Directory /api/v2")
    endpoint(doc, "GET", "/api/v2/lawyers/public", auth="No auth", query="?city=&practice=&lang=", note="Public verified lawyer directory.")
    endpoint(doc, "GET", "/api/v2/cas/public", auth="No auth", query="?city=", note="Public verified CA directory.")
    endpoint(doc, "GET", "/api/v2/lawyers", auth="Bearer JWT required", query="?city=&practice=&lang=", note="Full lawyer list including unverified.")

    # ── 11. Messaging ────────────────────────────────────────────────────────
    section(doc, "11. Messaging /api/v2/messages")
    body(doc, "Same routes mounted at /api/v2/messages, /api/v2/lawyer/messages, /api/v2/ca/messages. Source: routes/messageRoutes.js")

    endpoint(doc, "GET", "/api/v2/messages/threads", auth="Bearer JWT required", note="List threads + unreadCount.")
    endpoint(doc, "GET", "/api/v2/messages/threads/:id", auth="Bearer JWT required", note="Thread detail with messages.")
    endpoint(
        doc, "POST", "/api/v2/messages/threads", auth="Bearer JWT required",
        roles="client only",
        headers="Content-Type: application/json | Authorization: Bearer <token>",
        body='Body (JSON):\n{\n  "lawyerProfileId": 1,\n  "text": "Hello, I need consultation",\n  "attachments": []\n}',
        note="Start new thread with lawyer. lawyerUserId accepted as alias.",
    )
    endpoint(
        doc, "POST", "/api/v2/messages/threads/:id/messages", auth="Bearer JWT required",
        roles="client, lawyer, ca",
        headers="Content-Type: application/json | Authorization: Bearer <token>",
        body='Body (JSON):\n{\n  "text": "Follow-up message",\n  "attachments": []\n}',
        note="Send message in existing thread.",
    )
    endpoint(doc, "PATCH", "/api/v2/messages/threads/:id/read", auth="Bearer JWT required", note="Mark thread as read.")
    endpoint(doc, "GET", "/api/v2/messages/unread-count", auth="Bearer JWT required", note="Returns { unreadCount }.")

    # ── 12. Lawyer Dashboard ─────────────────────────────────────────────────
    section(doc, "12. Lawyer Dashboard /api/v2/lawyer")
    body(doc, "Source: routes/lawyerRoutes.js. Auth: Bearer JWT + lawyer role.")

    lawyer_routes = [
        ("GET", "/dashboard", "Dashboard summary KPIs"),
        ("GET", "/subscription", "Lawyer subscription info"),
        ("GET", "/notifications", "List notifications"),
        ("DELETE", "/notifications/:id", "Dismiss notification"),
        ("DELETE", "/notifications", "Clear all notifications"),
        ("GET", "/appointments", "Lawyer appointments"),
        ("PATCH", "/appointments/:appointmentId", "Update appointment status. Body: { status }"),
        ("GET", "/orders", "Assigned document orders"),
        ("POST", "/orders/:orderId/deliver", "Deliver order. multipart: document"),
        ("POST", "/orders/:orderId/esign", "Start e-sign session"),
        ("GET", "/vlo/subscribers", "VLO subscriber list"),
        ("GET", "/vlo/subscribers/:subscriberId/matters", "Matters for subscriber"),
        ("PATCH", "/vlo/matters/:matterId", "Update matter status"),
        ("POST", "/vlo/matters/:matterId/notes", "Add matter note. Body: { note }"),
        ("GET", "/clients", "Client list"),
        ("GET", "/clients/:clientId/history", "Client history (stub)"),
        ("GET", "/earnings", "Earnings summary"),
        ("GET", "/cases", "Case list"),
        ("POST", "/cases", "Create case"),
        ("PATCH", "/cases/:caseId", "Update case"),
        ("DELETE", "/cases/:caseId", "Delete case"),
        ("GET", "/profile", "Lawyer profile"),
        ("PATCH", "/profile", "Update profile fields"),
        ("POST", "/profile/photo", "Upload photo. multipart: photo"),
        ("GET", "/team", "Team members"),
        ("POST", "/team", "Add team member"),
        ("DELETE", "/team/:memberId", "Remove team member"),
        ("GET", "/lex/usage", "LEX AI usage statistics"),
        ("POST", "/lex/query", 'LEX chat. Body: { message, session_key }. Proxies to :8001'),
        ("POST", "/lexisnexis/connect", "LexisNexis integration placeholder"),
    ]
    for method, path, note in lawyer_routes:
        full = f"/api/v2/lawyer{path}"
        hdr = "Authorization: Bearer <token>"
        if method in ("POST", "PATCH"):
            hdr = "Content-Type: application/json | Authorization: Bearer <token>"
        endpoint(doc, method, full, auth="Bearer JWT + lawyer", headers=hdr, note=note)

    # ── 13. CA Dashboard ─────────────────────────────────────────────────────
    section(doc, "13. CA Dashboard /api/v2/ca")
    body(doc, "Source: routes/caRoutes.js. Auth: Bearer JWT + CA role.")

    ca_routes = [
        ("GET", "/dashboard", "CA dashboard summary"),
        ("GET", "/subscription", "CA subscription"),
        ("GET", "/notifications", "Notifications"),
        ("DELETE", "/notifications/:id", "Dismiss notification"),
        ("DELETE", "/notifications", "Clear notifications"),
        ("GET", "/compliance/deadlines", "Compliance deadlines"),
        ("GET", "/taxation/profiles", "Client tax profiles"),
        ("POST", "/taxation/profiles/:profileId/challans", "Upload challan. multipart: challan"),
        ("GET", "/orders", "CA service orders"),
        ("PATCH", "/orders/:orderId/milestone", "Update milestone. Body: { milestone }"),
        ("GET", "/documents", "CA documents"),
        ("POST", "/documents/:documentId/esign", "Trigger e-sign"),
        ("GET", "/retainers", "Retainer list"),
        ("GET", "/retainers/:retainerId/tasks", "Retainer tasks"),
        ("PATCH", "/retainers/tasks/:taskId", "Update task status"),
        ("GET", "/appointments", "CA appointments"),
        ("PATCH", "/appointments/:appointmentId", "Update appointment"),
        ("GET", "/profile", "CA profile"),
        ("PATCH", "/profile", "Update profile"),
        ("POST", "/profile/photo", "Upload photo. multipart: photo"),
        ("GET", "/team", "Team members"),
        ("POST", "/team", "Add team member"),
        ("DELETE", "/team/:memberId", "Remove team member"),
    ]
    for method, path, note in ca_routes:
        full = f"/api/v2/ca{path}"
        hdr = "Authorization: Bearer <token>"
        if method in ("POST", "PATCH"):
            hdr = "Content-Type: application/json | Authorization: Bearer <token>"
        endpoint(doc, method, full, auth="Bearer JWT + ca", headers=hdr, note=note)

    # ── 14. LEX AI ───────────────────────────────────────────────────────────
    section(doc, "14. LEX AI /api/v1/lex (Django)")
    body(doc, "Base :8001 or proxy http://localhost:3000/api/v1/lex. Source: lex_backend/lex_ai/views.py")

    endpoint(
        doc, "POST", "/api/v1/lex/chat/", auth="No auth",
        headers="Content-Type: application/json",
        body='Body (JSON):\n{\n  "message": "What is the limitation period for cheque bounce?",\n  "session_key": "session_1720556210000"\n}',
        note="RAG chat via Anthropic Claude. Returns { response, language, show_lawyer, shortcuts }.",
    )
    endpoint(doc, "GET", "/api/v1/lex/sessions/", auth="No auth", note="List all conversation sessions.")
    endpoint(doc, "GET", "/api/v1/lex/sessions/:sessionKey/", auth="No auth", note="Session message history.")

    # ── 15. WebSocket ────────────────────────────────────────────────────────
    section(doc, "15. WebSocket LEX Chat")
    endpoint(
        doc, "WS", "ws://localhost:3000/api/lex/ws", auth="No auth",
        body='Send (JSON):\n{\n  "query": "What is Article 199?",\n  "session_key": "ws_1720556210000"\n}',
        note='Receive: { text, shortcuts }. Proxies to LEX AI :8001. Source: server.js',
    )

    # ── 16. Static Files ─────────────────────────────────────────────────────
    section(doc, "16. Static Files")
    endpoint(doc, "GET", "/uploads/*", auth="No auth", note="Uploaded files (profile photos, documents). Source: server.js")

    # ── 17. Postman / Integration Tips ───────────────────────────────────────
    section(doc, "17. Postman / Integration Tips")
    tips = """1. Use http://localhost:3000/api/auth/* for all auth (single tunnel via ngrok)
2. Login → copy token → set Header: Authorization: Bearer <token>
3. Signup OTP flow: send-otp → verify-otp → register with verificationToken
4. switch-role returns new token — replace stored token after role switch
5. File uploads: use form-data (multipart), not raw JSON
6. Lawyer/CA dashboards require verified profile (admin approval)
7. Main API dev mock: token mock-jwt-token-* + X-Client-Role: CorporateClient
8. LEX WebSocket for real-time chat; REST /api/v1/lex/chat/ for HTTP clients

Example login success:
{
  "token": "eyJhbGciOi...",
  "user": {
    "id": 1,
    "name": "Ali Khan",
    "email": "ali@example.com",
    "role": "client",
    "roles": ["CorporateClient"],
    "profile": { "profileCompleted": true, ... }
  }
}

Example OTP verify success:
{
  "ok": true,
  "email": "ali@example.com",
  "verificationToken": "a1b2c3...",
  "message": "Email verified successfully."
}

Example role switch blocked (409):
{
  "error": "Lawyer verification is still pending",
  "code": "PENDING_VERIFICATION",
  "pendingRole": "Lawyer"
}"""
    for line in tips.split("\n"):
        mono(doc, line)

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— End of Nexus Lexis API Documentation —")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    return doc


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(str(OUT))
    print(f"Created: {OUT}")


if __name__ == "__main__":
    main()
