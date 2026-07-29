"""Convert docs/BACKEND-API-KEYS-REFERENCE.md to a Word document."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs" / "BACKEND-API-KEYS-REFERENCE.md"
OUT_PATH = ROOT / "docs" / "BACKEND-API-KEYS-REFERENCE.docx"

DOC_TITLE = "Nexus Lexis — Backend API Keys & Secrets Reference"
DOC_ID = "NL-DOC-SEC-001"
DOC_VERSION = "1.0"
DOC_DATE = "2026-07-28"
DOC_CLASS = "INTERNAL — CONFIDENTIAL"


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = f"{DOC_TITLE}  |  {DOC_ID}  |  v{DOC_VERSION}  |  {DOC_CLASS}"
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in hp.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.font.name = "Calibri"

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = (
        f"© Nexus Lexis · nexuslexis.law · contact@nexuslexis.law  |  "
        f"Last updated: {DOC_DATE}  |  Page "
    )
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in fp.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.font.name = "Calibri"

    # Page number field
    run = fp.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_title_page(doc: Document, meta: dict[str, str]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NEXUS LEXIS PLATFORM")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x16, 0x21, 0x3E)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Backend API Keys, Secrets & Authentication Headers")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x16, 0x21, 0x3E)

    doc.add_paragraph()

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    fields = [
        ("Document ID", meta.get("Document ID", DOC_ID)),
        ("Version", meta.get("Version", DOC_VERSION)),
        ("Last Updated", meta.get("Last Updated", DOC_DATE)),
        ("Classification", meta.get("Classification", DOC_CLASS)),
        ("Owner", meta.get("Owner", "Nexus Lexis Engineering")),
        ("Applies To", meta.get("Applies To", "auth_backend · Main API · lex_backend")),
    ]
    for label, value in fields:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value
        set_cell_shading(row[0], "E8ECF4")
        for cell in row:
            for para in cell.paragraphs:
                for r in para.runs:
                    r.font.name = "Calibri"
                    r.font.size = Pt(10)

    doc.add_page_break()


def parse_front_matter(lines: list[str]) -> tuple[dict[str, str], list[str]]:
    meta: dict[str, str] = {}
    body: list[str] = []
    in_meta = False

    for line in lines:
        if line.strip() == "---":
            in_meta = not in_meta
            continue
        if in_meta:
            if ":" in line:
                key, val = line.split(":", 1)
                meta[key.strip()] = val.strip()
        else:
            body.append(line)
    return meta, body


def add_rich_text(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


def add_table_from_markdown(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            cell = table.rows[r_idx].cells[c_idx]
            text = row[c_idx] if c_idx < len(row) else ""
            cell.text = ""
            para = cell.paragraphs[0]
            add_rich_text(para, text.strip())
            for run in para.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(9)
            if r_idx == 0:
                set_cell_shading(cell, "D4E4F7")
                for run in para.runs:
                    run.bold = True


def convert_markdown_to_docx(md_text: str) -> Document:
    meta, body_lines = parse_front_matter(md_text.splitlines())
    doc = Document()
    add_header_footer(doc)
    add_title_page(doc, meta)

    in_code = False
    code_lines: list[str] = []
    table_rows: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_rows
        if table_rows:
            add_table_from_markdown(doc, table_rows)
            table_rows = []
            doc.add_paragraph()

    for raw in body_lines:
        line = raw.rstrip("\n")

        if in_code:
            if line.strip().startswith("```"):
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                in_code = False
                code_lines = []
            else:
                code_lines.append(line)
            continue

        if line.strip().startswith("```"):
            flush_table()
            in_code = True
            code_lines = []
            continue

        if line.strip().startswith("|") and "|" in line.strip()[1:]:
            if re.match(r"^\|\s*[-: ]+\|", line.strip()):
                continue
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            table_rows.append(cells)
            continue

        flush_table()

        if line.strip() == "---":
            doc.add_paragraph()
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
            continue

        if line.startswith("> "):
            p = doc.add_paragraph(style="Intense Quote")
            add_rich_text(p, line[2:].strip())
            continue

        if not line.strip():
            continue

        p = doc.add_paragraph()
        add_rich_text(p, line.strip())

    flush_table()
    return doc


def main() -> None:
    md_text = MD_PATH.read_text(encoding="utf-8")
    doc = convert_markdown_to_docx(md_text)
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"Created: {OUT_PATH}")


if __name__ == "__main__":
    main()
